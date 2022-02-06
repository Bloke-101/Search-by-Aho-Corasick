from io import StringIO
from docx import Document
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfpage import PDFPage
import os
import shutil

PATH_TO_PATTERNS = "\path\to\patterns.txt"


class Trie:
    def __init__(self, is_root=False):
        self.next = {}
        self.is_end = False
        self.suf_link = -1
        self.short_suf_link = -1
        self.is_root = is_root


def insert(node, pat):
    cur = node
    for i in pat:
        if not cur.next.get(i):
            cur.next[i] = Trie()
        cur = cur.next[i]
    cur.is_end = True


def bfs(node):
    queue = []
    node.suf_link = node
    node.short_suf_link = None
    for key, value in node.next.items():
        value.suf_link = node
        value.short_suf_link = None
        queue.append(value)
    while queue:
        v = queue.pop(0)
        for key, value in v.next.items():
            par_ch = key
            child = value
            par_suf = v.suf_link
            while not (par_suf.is_root or par_suf.next.get(par_ch)):
                par_suf = par_suf.suf_link
            next_suf = par_suf.next[par_ch] if par_suf.next.get(par_ch) else par_suf
            child.suf_link = next_suf
            v.short_suf_link = get_short_suf_link(v)
            queue.append(child)


def get_short_suf_link(node):
    if node.short_suf_link == -1:
        if node.suf_link.is_end:
            node.short_suf_link = node.suf_link
        elif node.is_end:
            node.short_suf_link = node
        elif node.suf_link.is_root:
            node.short_suf_link = None
        else:
            node.short_suf_link = get_short_suf_link(node.suf_link)
    return node.short_suf_link


def go(node, char):
    cur = node
    while not (cur.is_root or cur.next.get(char)):
        cur = cur.suf_link
    if cur.next.get(char):
        return cur.next[char]
    return cur


def find_all_patterns(node, string):
    cur = node
    count = 0
    for i in string:
        cur = go(cur, i)
        if cur.is_end:
            count += 1
        else:
            temp = cur.short_suf_link
            while temp:
                count += 1
                temp = temp.short_suf_link
    return count


def read_txt(path):
    with open(path) as file:
        txt_text = file.read()
    return txt_text


def read_docx(path):
    docx_text = []
    doc = Document(path)
    for par in doc.paragraphs:
        docx_text.append(par.text)
    tables = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tables.append(cell.text)
    docx_text.extend(tables)
    return " ".join(docx_text)


def read_pdf(path):
    resource_manger = PDFResourceManager()
    file_handle = StringIO()
    converter = TextConverter(resource_manger, file_handle)
    page_interpreter = PDFPageInterpreter(resource_manger, converter)
    with open(path, "rb") as f:
        for page in PDFPage.get_pages(f,
                                      caching=True,
                                      check_extractable=True):
            page_interpreter.process_page(page)
            pdf_text = file_handle.getvalue()
    converter.close()
    file_handle.close()
    return pdf_text


def get_text(path_file):
    full_name = os.path.basename(path_file)
    extension = os.path.splitext(full_name)[1]
    if extension == ".txt":
        return read_txt(path_file)
    elif extension == ".docx":
        return read_docx(path_file)
    elif extension == ".pdf":
        return read_pdf(path_file)
    else:
        return "None"


def add_pattern(pat):
    with open(PATH_TO_PATTERNS, mode="a") as file:
        file.write(f"\n{pat}")


def find_word_count(string):
    string = string.replace("\n", " ")
    string = string.lower()
    string = string.strip()
    words = string.split()
    return len(words)


root = Trie(is_root=True)
patterns = read_txt(PATH_TO_PATTERNS).lower().split("\n")
for pattern in patterns:
    insert(root, pattern)
bfs(root)
while True:
    print("Enter 1 to add new pattern\nEnter 2 to sort the files\nEnter 3 to exit")
    while True:
        try:
            choice = int(input())
            if choice != 1 and choice != 2 and choice != 3:
                continue
            break
        except ValueError:
            print("Invalid input. Please, enter the correct value.")
    if choice == 1:
        add_pattern(input("Enter new pattern: "))
    elif choice == 2:
        while True:
            try:
                dir_path = input("Enter path to directory: ")
                folders = os.listdir(dir_path)
                break
            except FileNotFoundError:
                print("File not found. Enter the path again.")
        if folders:
            if not os.path.isdir(f"{dir_path}\\spam"):
                os.mkdir(f"{dir_path}\\spam")
            for folder in folders:
                text = get_text(f"{dir_path}\\{folder}")
                if text != "None":
                    word_count = find_word_count(text)
                    res = find_all_patterns(root, text.lower())
                    if word_count != 0 and res / word_count >= 0.1:
                        shutil.move(f"{dir_path}\\{folder}", f"{dir_path}\\spam")
    else:
        break
